import os
import subprocess
import jinja2
from datetime import date
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- HELPER: ESCAPE LATEX ---
def escape_latex(text):
    if not isinstance(text, str):
        return text
    chars = {
        '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_',
        '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}', '^': r'\^{}',
        '\\': r'\textbackslash{}', '<': r'\textless{}', '>': r'\textgreater{}',
        "'": r"'",  # Use proper LaTeX quote
    }
    return "".join(chars.get(c, c) for c in text)

# --- PDF GENERATOR (Improved Error Handling) ---
def generate_resume_pdf(data, template_name="modern", output_dir="output"):
    # 1. Setup Jinja2
    template_loader = jinja2.FileSystemLoader(searchpath="./assets/templates")
    latex_jinja_env = jinja2.Environment(
        loader=template_loader,
        block_start_string='\\BLOCK{', block_end_string='}',
        variable_start_string='\\VAR{', variable_end_string='}',
        comment_start_string='\\#{', comment_end_string='}',
        line_statement_prefix='%%', line_comment_prefix='%#',
        trim_blocks=True, autoescape=False,
    )
    
    template_file = f"{template_name}.tex"
    try:
        template = latex_jinja_env.get_template(template_file)
    except jinja2.TemplateNotFound:
        print(f"Template {template_file} not found. Falling back to modern.tex")
        template = latex_jinja_env.get_template('modern.tex')

    # 3. Clean Data for LaTeX
    clean_data = {}
    def clean_list(item_list):
        cleaned = []
        for item in item_list:
            new_item = {}
            for k, v in item.items():
                if isinstance(v, list):
                    new_item[k] = [escape_latex(b) for b in v]
                else:
                    new_item[k] = escape_latex(v)
            cleaned.append(new_item)
        return cleaned

    for field in ['name', 'email', 'phone', 'linkedin', 'github', 'summary', 'website']:
        clean_data[field] = escape_latex(data.get(field, ''))

    clean_data['experience'] = clean_list(data.get('experience', []))
    clean_data['education'] = clean_list(data.get('education', []))
    clean_data['projects'] = clean_list(data.get('projects', []))
    
    clean_data['skills'] = []
    for idx, s in enumerate(data.get('skills', [])):
        try:
            # Direct access to avoid .get() gotcha where dict.get('items') returns the method
            items_value = s['items']
            
            # Convert items to string if it's not already
            if isinstance(items_value, str):
                items_str = items_value
            elif isinstance(items_value, list):
                items_str = ', '.join(str(item) for item in items_value)
            elif isinstance(items_value, dict):
                # If it's a dict, try to extract meaningful values
                items_str = ', '.join(str(v) for v in items_value.values() if v)
            elif callable(items_value):
                # Skip if it's a method (the dict.items gotcha)
                print(f"⚠️ Warning: Skill {idx} 'items' is callable, skipping")
                continue
            else:
                items_str = str(items_value)
            
            # Only add if we have valid data
            if items_str and items_str.strip():
                clean_data['skills'].append({
                    'category': escape_latex(str(s.get('category', ''))),
                    'items': escape_latex(items_str)
                })
            else:
                print(f"⚠️ Warning: Skill {idx} has empty items, skipping")
                
        except KeyError as e:
            print(f"❌ Skill {idx} missing key: {e}")
            continue
        except Exception as e:
            print(f"❌ Error processing skill {idx}: {e}")
            continue
    
    # If no valid skills, add a placeholder
    if not clean_data['skills']:
        print("⚠️ No valid skills found, using placeholder")
        clean_data['skills'] = [{
            'category': 'Skills',
            'items': 'Please update your skills section'
        }]
    
    # Validate and ensure all sections have at least one entry
    if not clean_data.get('experience'):
        clean_data['experience'] = [{
            'title': 'Your Job Title',
            'company': 'Company Name',
            'dates': '2020 - Present',
            'bullets': ['Add your experience details']
        }]
    
    if not clean_data.get('education'):
        clean_data['education'] = [{
            'school': 'Your University',
            'degree': 'Your Degree',
            'year': '2020',
            'gpa': '3.5'
        }]
    
    if not clean_data.get('projects'):
        clean_data['projects'] = [{
            'name': 'Your Project',
            'description': 'Project description',
            'link': ''
        }]
    
    # Ensure all experience entries have bullets
    for exp in clean_data.get('experience', []):
        if not exp.get('bullets') or len(exp.get('bullets', [])) == 0:
            exp['bullets'] = ['Responsibilities and achievements']
    
    # Ensure all required fields exist with defaults
    for exp in clean_data.get('experience', []):
        exp.setdefault('title', 'Position')
        exp.setdefault('company', 'Company')
        exp.setdefault('dates', '2020')
    
    for edu in clean_data.get('education', []):
        edu.setdefault('school', 'University')
        edu.setdefault('degree', 'Degree')
        edu.setdefault('year', '2020')
        edu.setdefault('gpa', '3.0')
    
    for proj in clean_data.get('projects', []):
        proj.setdefault('name', 'Project')
        proj.setdefault('description', 'Description')
        proj.setdefault('link', '')




    rendered_tex = template.render(**clean_data, today=date.today().strftime("%B %Y"))

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    tex_path = os.path.join(output_dir, "resume.tex")
    with open(tex_path, "w", encoding='utf-8') as f:
        f.write(rendered_tex)

    # Compile LaTeX with proper error handling
    try:
        result = subprocess.run(
            ['pdflatex', '-interaction=nonstopmode', 
             f'-output-directory={output_dir}', tex_path],
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode != 0:
            # Extract the actual error from LaTeX output
            error_lines = []
            for line in result.stdout.split('\n'):
                if line.startswith('!') or 'Error' in line or 'error' in line:
                    error_lines.append(line)
            
            error_msg = f"LaTeX compilation failed with return code {result.returncode}"
            if error_lines:
                error_msg += f"\n\nLaTeX Errors:\n" + "\n".join(error_lines[:10])
            else:
                error_msg += f"\n\nOutput (last 500 chars):\n{result.stdout[-500:]}"
            
            raise Exception(error_msg)
            
        pdf_path = os.path.join(output_dir, "resume.pdf")
        if not os.path.exists(pdf_path):
            raise Exception("PDF was not generated despite successful compilation")
            
        return pdf_path
        
    except FileNotFoundError:
        raise Exception(
            "pdflatex not found. Please install LaTeX:\n"
            "- Windows: Install MiKTeX from https://miktex.org/\n"
            "- Mac: Install MacTeX from https://www.tug.org/mactex/\n"
            "- Linux: sudo apt-get install texlive-full"
        )
    except subprocess.TimeoutExpired:
        raise Exception("LaTeX compilation timed out after 30 seconds")


# --- HELPER: BORDER FOR WORD ---
def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

# --- UPGRADED WORD GENERATOR (WITH NARROW MARGINS) ---
def generate_resume_docx(data, output_dir="output"):
    doc = DocxDocument()
    
    # 1. SET NARROW MARGINS (0.5 inches)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # 2. Set Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # 3. HEADER
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data.get('name', 'Name'))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.name = 'Times New Roman'
    
    contact_info = []
    if data.get('email'): contact_info.append(data.get('email'))
    if data.get('phone'): contact_info.append(data.get('phone'))
    if data.get('linkedin'): contact_info.append(f"LinkedIn: {data.get('linkedin')}")
    if data.get('github'): contact_info.append(f"GitHub: {data.get('github')}")
    
    contact_para = doc.add_paragraph(" | ".join(contact_info))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(10)

    # Helper for Titles
    def add_section_header(title):
        p = doc.add_heading(title.upper(), level=1)
        run = p.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        add_bottom_border(p)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    # 4. SUMMARY
    add_section_header('Professional Summary')
    doc.add_paragraph(data.get('summary', ''))

    # 5. EXPERIENCE (Adjusted Table Widths for Narrow Layout)
    add_section_header('Experience')
    for job in data.get('experience', []):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Wider columns because we have more space now!
        # Total width = ~7.5 inches available
        cell_1 = table.cell(0, 0)
        cell_1.width = Inches(5.5)  # Title area
        p1 = cell_1.paragraphs[0]
        r1 = p1.add_run(f"{job.get('title', '')}")
        r1.bold = True
        p1.add_run(f" | {job.get('company', '')}").italic = True
        
        cell_2 = table.cell(0, 1)
        cell_2.width = Inches(2.0)  # Date area
        p2 = cell_2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.add_run(job.get('dates', '')).italic = True
        
        if 'bullets' in job:
            for bullet in job['bullets']:
                b_para = doc.add_paragraph(bullet, style='List Bullet')
                b_para.paragraph_format.space_after = Pt(0)

    # 6. PROJECTS
    if data.get('projects'):
        add_section_header('Projects')
        for proj in data.get('projects', []):
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            
            cell_1 = table.cell(0, 0)
            cell_1.width = Inches(5.5)
            p1 = cell_1.paragraphs[0]
            p1.add_run(proj.get('name', '')).bold = True
            
            cell_2 = table.cell(0, 1)
            cell_2.width = Inches(2.0)
            p2 = cell_2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if proj.get('link'):
                p2.add_run(proj.get('link', '')).italic = True

            if proj.get('description'):
                b_para = doc.add_paragraph(proj.get('description'), style='List Bullet')
                b_para.paragraph_format.space_after = Pt(2)

    # 7. EDUCATION
    if data.get('education'):
        add_section_header('Education')
        for edu in data.get('education', []):
            table = doc.add_table(rows=1, cols=2)
            c1 = table.cell(0,0)
            c1.width = Inches(5.5)
            c1.paragraphs[0].add_run(edu.get('school', '')).bold = True
            
            c2 = table.cell(0,1)
            c2.width = Inches(2.0)
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            c2.paragraphs[0].add_run(edu.get('year', ''))
            
            p = doc.add_paragraph()
            p.add_run(edu.get('degree', '')).italic = True
            p.add_run(f" | GPA: {edu.get('gpa', '')}")
            p.paragraph_format.space_after = Pt(6)

    # 8. SKILLS
    add_section_header('Technical Skills')
    for skill in data.get('skills', []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"{skill.get('category')}: ").bold = True
        p.add_run(skill.get('items'))

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    docx_path = os.path.join(output_dir, "Optimized_Resume.docx")
    doc.save(docx_path)
    return docx_path